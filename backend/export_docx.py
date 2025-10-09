#!/usr/bin/env python3
"""
Export DOCX à partir du JSON de contenu du CV
Utilise python-docx. Entrée: fichier JSON (même structure que data/*.json)
"""
import argparse
import json
from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def load_json(path: Path):
    with path.open("r", encoding="utf-8") as f:
        return json.load(f)


def add_heading(doc: Document, text: str, level: int = 1):
    p = doc.add_paragraph()
    run = p.add_run(text)
    if level == 1:
        run.font.size = Pt(16)
        run.bold = True
    elif level == 2:
        run.font.size = Pt(13)
        run.bold = True
    else:
        run.font.size = Pt(12)
        run.bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT


def add_bullets(doc: Document, items):
    for it in items:
        doc.add_paragraph(str(it), style="List Bullet")


def export_docx(data_json_path: Path, out_docx_path: Path):
    data = load_json(data_json_path)
    d_profile = data.get("profile", {})
    d_contacts = (d_profile or {}).get("contacts", {})

    doc = Document()

    # En-tête
    name = (d_profile or {}).get("name") or ""
    if name:
        h = doc.add_heading(level=0)
        h_run = h.add_run(name)
        h_run.font.size = Pt(20)
        h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title = (d_profile or {}).get("title")
    if title:
        p = doc.add_paragraph(title)
        p.runs[0].font.size = Pt(12)

    # Contacts
    contacts_line = " | ".join(
        x for x in [d_contacts.get("email"), d_contacts.get("phone"), d_contacts.get("location")] if x
    )
    if contacts_line:
        p = doc.add_paragraph(contacts_line)
        p.runs[0].font.size = Pt(10)

    # Résumé
    summary = data.get("summary")
    if summary:
        add_heading(doc, "Résumé", level=1)
        doc.add_paragraph(summary)

    # Compétences
    skills = data.get("skills", {})
    groups = skills.get("groups") if isinstance(skills, dict) else None
    if groups:
        add_heading(doc, "Compétences", level=1)
        for grp in groups:
            label = grp.get("label")
            items = grp.get("items", [])
            if label:
                add_heading(doc, label, level=2)
            add_bullets(doc, items)

    # Expérience
    xp = data.get("experience", [])
    if xp:
        add_heading(doc, "Expériences", level=1)
        for item in xp:
            title_line = " - ".join([x for x in [item.get("role"), item.get("company")] if x])
            if title_line:
                add_heading(doc, title_line, level=2)
            date_line = " - ".join([x for x in [item.get("start"), item.get("end")] if x])
            if date_line:
                doc.add_paragraph(date_line)
            bullets = item.get("bullets", [])
            add_bullets(doc, bullets)

    # Formation
    edu = data.get("education", [])
    if edu:
        add_heading(doc, "Formation", level=1)
        for e in edu:
            title_line = " - ".join([x for x in [e.get("degree"), e.get("school")] if x])
            if title_line:
                add_heading(doc, title_line, level=2)
            year = e.get("year")
            if year:
                doc.add_paragraph(str(year))

    doc.save(str(out_docx_path))
    return out_docx_path


def main():
    parser = argparse.ArgumentParser(description="Export DOCX from resume JSON")
    parser.add_argument("--data", required=True, help="Path to data JSON (e.g., data/resume_marlyse.json)")
    parser.add_argument("--out", required=False, default="CV.docx", help="Output DOCX path")
    args = parser.parse_args()

    out = export_docx(Path(args.data), Path(args.out))
    print(f"DOCX généré: {out}")


if __name__ == "__main__":
    main()
